"""
Document ingestion pipeline for the Nursing AI Assistant.

Mirrors information_vectorization_v4.ipynb so that documents uploaded through the
app are chunked, enriched and embedded *exactly* like the original knowledge base
and land in the same Chroma collection.

Flow per uploaded file
-----------------------
    1. save        — store the original file under data/uploads/<source>/
    2. convert     — turn it into a markdown file (+ extracted image assets)
    3. parse       — split the markdown into text / table / image chunks
    4. enrich      — describe images (vision model) and summarise tables (LLM)
                     via Ollama. Skipped gracefully if Ollama is unavailable.
    5. split       — recursive character splitting (1400 / 200, like the notebook)
    6. keywords    — short keyword list per chunk (LLM, optional)
    7. embed       — add to the Chroma collection, overwriting any prior chunks
                     that share the same `source`.

The public entry point is `ingest_file(...)`. Everything reports progress through
an optional `log` callback so the Streamlit page can show what is happening.
"""

from __future__ import annotations

import os
import re
import shutil
from pathlib import Path
from typing import Callable, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma

from config import (
    VECTOR_DB_PATH,
    COLLECTION_NAME,
    LLM_MODEL,
    VISION_MODEL,
    INGEST_CHUNK_SIZE,
    INGEST_CHUNK_OVERLAP,
    INGEST_MAX_IMAGE_RESOLUTION,
    INGEST_MAX_IMAGE_TOKENS,
    UPLOADS_DIR,
)

Logger = Callable[[str], None]


def _noop(_msg: str) -> None:
    pass


# ===========================================================================
# Ollama availability  (the enrichment steps degrade gracefully without it)
# ===========================================================================
def ollama_status() -> dict:
    """Return what Ollama can do right now: {'up', 'models', 'has_llm', 'has_vision'}."""
    info = {"up": False, "models": [], "has_llm": False, "has_vision": False}
    try:
        import ollama

        listed = ollama.list()
        names = [m.get("model", m.get("name", "")) for m in listed.get("models", [])]
        info["up"] = True
        info["models"] = names

        def _present(target: str) -> bool:
            base = target.split(":")[0]
            return any(n == target or n.split(":")[0] == base for n in names)

        info["has_llm"] = _present(LLM_MODEL)
        info["has_vision"] = _present(VISION_MODEL)
    except Exception:
        pass
    return info


def sanitize_source_name(name: str) -> str:
    """Turn a filename into a safe, human-readable source/document name."""
    stem = Path(name).stem.strip()
    stem = re.sub(r"[\\/:*?\"<>|]+", "_", stem)
    return stem or "document"


# ===========================================================================
# STEP 2 — Convert an uploaded file into  content.md  (+ assets/)
# ===========================================================================
def _write_markdown(source_dir: Path, markdown: str) -> Path:
    content_path = source_dir / "content.md"
    content_path.write_text(markdown, encoding="utf-8")
    return content_path


def _convert_txt(original: Path, source_dir: Path, log: Logger) -> Path:
    text = original.read_text(encoding="utf-8", errors="ignore")
    log("Read plain text.")
    return _write_markdown(source_dir, text)


def _convert_md(original: Path, source_dir: Path, log: Logger) -> Path:
    text = original.read_text(encoding="utf-8", errors="ignore")
    # Bring along a sibling assets folder if the author shipped one.
    sib = original.parent / "assets"
    if sib.is_dir():
        shutil.copytree(sib, source_dir / "assets", dirs_exist_ok=True)
    log("Copied markdown as-is.")
    return _write_markdown(source_dir, text)


def _convert_pdf(original: Path, source_dir: Path, log: Logger) -> Path:
    """Extract headings (font-size heuristic), text, tables and images from a PDF."""
    import fitz  # PyMuPDF

    assets_dir = source_dir / "assets"
    assets_dir.mkdir(exist_ok=True)

    doc = fitz.open(str(original))
    out_lines: list[str] = []
    img_count = 0
    tbl_count = 0

    # Establish a "body" font size to detect headings.
    sizes: list[float] = []
    for page in doc:
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    sizes.append(round(span["size"], 1))
    body_size = max(set(sizes), key=sizes.count) if sizes else 11.0

    def _inside(table_bboxes, bbox) -> bool:
        cx = (bbox[0] + bbox[2]) / 2
        cy = (bbox[1] + bbox[3]) / 2
        return any(tb[0] <= cx <= tb[2] and tb[1] <= cy <= tb[3] for tb in table_bboxes)

    for pno, page in enumerate(doc, 1):
        # --- tables first (so we can skip their text in the body) ---
        table_bboxes = []
        try:
            tables = page.find_tables()
            for tbl in tables.tables:
                table_bboxes.append(tbl.bbox)
                md = tbl.to_markdown() or ""
                if md.strip():
                    tbl_count += 1
                    out_lines.append("")
                    out_lines.append(md.strip())
                    out_lines.append("")
        except Exception:
            pass

        # --- text, with a light heading heuristic ---
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                if _inside(table_bboxes, line["bbox"]):
                    continue
                spans = line.get("spans", [])
                text = "".join(s["text"] for s in spans).strip()
                if not text:
                    continue
                max_size = max((s["size"] for s in spans), default=body_size)
                is_bold = any("bold" in s.get("font", "").lower() for s in spans)
                if (max_size >= body_size * 1.25 or (is_bold and max_size >= body_size * 1.1)) and len(text) <= 120:
                    level = "##" if max_size >= body_size * 1.6 else "###"
                    out_lines.append(f"\n{level} {text}\n")
                else:
                    out_lines.append(text)

        # --- images ---
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.width < 60 or pix.height < 60:  # skip icons / rules
                    pix = None
                    continue
                if pix.n - pix.alpha >= 4:  # CMYK / other -> RGB
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                img_count += 1
                fname = f"page_{pno}_img_{img_count}.png"
                pix.save(str(assets_dir / fname))
                pix = None
                out_lines.append(f"\n![figure](assets/{fname})\n")
            except Exception:
                continue

    doc.close()
    log(f"Extracted PDF: {tbl_count} table(s), {img_count} image(s).")
    return _write_markdown(source_dir, "\n".join(out_lines))


def _convert_docx(original: Path, source_dir: Path, log: Logger) -> Path:
    """Extract headings, paragraphs, tables and images from a .docx in reading order."""
    import docx
    from docx.document import Document as _Doc
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    assets_dir = source_dir / "assets"
    assets_dir.mkdir(exist_ok=True)

    document = docx.Document(str(original))
    out_lines: list[str] = []
    tbl_count = 0

    def _iter_block_items(parent):
        """Yield paragraphs and tables in document order."""
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P

        body = parent.element.body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _table_to_md(table: Table) -> str:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            return ""
        header = rows[0]
        md = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
        for r in rows[1:]:
            md.append("| " + " | ".join(r) + " |")
        return "\n".join(md)

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            style = (block.style.name or "").lower() if block.style else ""
            if style.startswith("heading"):
                m = re.search(r"(\d+)", style)
                level = "#" * min(int(m.group(1)) + 1, 6) if m else "##"
                out_lines.append(f"\n{level} {text}\n")
            elif style == "title":
                out_lines.append(f"\n# {text}\n")
            else:
                out_lines.append(text)
        elif isinstance(block, Table):
            md = _table_to_md(block)
            if md:
                tbl_count += 1
                out_lines.append("")
                out_lines.append(md)
                out_lines.append("")

    # Images live in the package relationships; extract and append them.
    img_count = 0
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
                ext = (rel.target_part.partname.ext or "png").lstrip(".")
                img_count += 1
                fname = f"image_{img_count}.{ext}"
                (assets_dir / fname).write_bytes(blob)
                out_lines.append(f"\n![figure](assets/{fname})\n")
            except Exception:
                continue

    log(f"Extracted DOCX: {tbl_count} table(s), {img_count} image(s).")
    return _write_markdown(source_dir, "\n".join(out_lines))


_CONVERTERS = {
    ".pdf": _convert_pdf,
    ".docx": _convert_docx,
    ".md": _convert_md,
    ".markdown": _convert_md,
    ".txt": _convert_txt,
}


def convert_to_markdown(original: Path, source_dir: Path, log: Logger = _noop) -> Path:
    ext = original.suffix.lower()
    converter = _CONVERTERS.get(ext)
    if not converter:
        raise ValueError(f"Unsupported file type: {ext}")
    return converter(original, source_dir, log)


# ===========================================================================
# STEP 3 — Parse markdown into text / table / image chunks
#          (ported verbatim from information_vectorization_v4.ipynb)
# ===========================================================================
def parse_markdown_to_chunks(file_path: str, source_name: str) -> list[dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    chunks: list[dict] = []
    header_stack: dict[int, str] = {}
    current_hierarchy = "Root"
    current_content: list[str] = []
    current_type = "only_text"

    def commit_chunk():
        nonlocal current_content, current_type
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                chunks.append({
                    "content": text,
                    "metadata": {"Hierarchy": current_hierarchy, "source": source_name, "asset": "None"},
                    "category": current_type,
                })
            current_content = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        header_match = re.match(r"^(#+)\s+(.*)", stripped)
        if header_match:
            commit_chunk()
            level = len(header_match.group(1))
            title = header_match.group(2)
            header_stack[level] = title
            for k in list(header_stack.keys()):
                if k > level:
                    del header_stack[k]
            current_hierarchy = " > ".join(header_stack[k] for k in sorted(header_stack.keys()))
            current_type = "only_text"
            continue

        if stripped.startswith("![") and "](" in stripped:
            commit_chunk()
            current_type = "contains_image"
            current_content.append(stripped)
            commit_chunk()
            current_type = "only_text"
            continue

        if "|" in stripped and (stripped.startswith("|") or current_type == "contains_table"):
            if current_type != "contains_table":
                commit_chunk()
                current_type = "contains_table"
            current_content.append(stripped)
            continue

        if current_type == "contains_table":
            commit_chunk()
            current_type = "only_text"

        if stripped != "RESTRICTED":
            current_content.append(stripped)

    commit_chunk()
    return chunks


# ===========================================================================
# STEP 4 — Enrichment (images + tables) via Ollama, with graceful fallback
# ===========================================================================
def clean_markdown_table(table_content: str) -> str:
    cleaned = re.sub(r"<br\s*/?>", " ", table_content, flags=re.IGNORECASE)
    cleaned = re.sub(r"\(\?\)", "", cleaned)
    cleaned = re.sub(r"(?<!\w)\?(?!\w)", "", cleaned)
    cleaned = re.sub(r" {2,}", " ", cleaned)
    cleaned = re.sub(r"(?m)^\|\s*(\|\s*)+\n?", "", cleaned)
    return cleaned.strip()


def summarize_table_with_qwen(table_content: str) -> str:
    system_prompt = (
        "You are an expert clinical assistant. Analyze the Markdown table "
        "and summarize it into clear, well-structured paragraphs.\n"
        "Rules:\n"
        "1. Maintain all clinical concepts, categories, levels and comparisons.\n"
        "2. Explain relationships between columns clearly.\n"
        "3. Output only the summarized text."
    )
    import ollama

    response = ollama.chat(
        model=LLM_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Summarize this table:\n\n{table_content}"},
        ],
    )
    return response["message"]["content"].strip()


def summarize_image_with_vision(chunk: dict, base_dir: Path) -> str:
    """Describe the image referenced in the chunk. base_dir is the source folder."""
    import ollama
    from PIL import Image

    content = chunk["content"]
    metadata = chunk["metadata"]
    match = re.search(r"!\[.*?\]\((.*?)\)", content)
    if not match:
        return content

    rel_path = match.group(1)
    metadata["asset"] = Path(rel_path).name
    original_image_path = (base_dir / rel_path)
    if not original_image_path.exists():
        return content

    resolution = INGEST_MAX_IMAGE_RESOLUTION
    max_tokens = INGEST_MAX_IMAGE_TOKENS
    res_str = f"{resolution[0]}x{resolution[1]}"
    optimized = original_image_path.parent / f"opt_{res_str}_{original_image_path.name}"
    try:
        if not optimized.exists():
            with Image.open(original_image_path) as img:
                img.thumbnail(resolution)
                if img.mode != "RGB":
                    img = img.convert("RGB")
                img.save(optimized, "JPEG", quality=85)
        image_to_process = str(optimized)
    except Exception:
        image_to_process = str(original_image_path)

    hierarchy = metadata.get("Hierarchy", "Unknown Section")
    word_limit = int(max_tokens * 0.75)
    system_prompt = (
        f"You are an expert clinical assistant.\n"
        f"Describe the clinical content, diagram, workflow or key elements shown in the image.\n"
        f"CRITICAL RULE: Limit your response to strictly UNDER {word_limit} words. Be concise.\n"
        f"Context Section: [{hierarchy}]"
    )
    response = ollama.chat(
        model=VISION_MODEL,
        messages=[{"role": "user", "content": system_prompt, "images": [image_to_process]}],
        options={"num_predict": max_tokens, "temperature": 0.2},
    )
    return f"Image Asset ({metadata['asset']}) Description: {response['message']['content'].strip()}"


def extract_keywords_with_llm(text: str) -> str:
    import ollama

    prompt = (
        "Extract 5-10 clinical keywords from the text below as a comma-separated list. "
        "Focus on psychiatric and nursing concepts, assessments and terminology.\n\n"
        f"Text: {text}"
    )
    response = ollama.generate(model=LLM_MODEL, prompt=prompt, options={"temperature": 0.0})
    return response["response"].strip()


# ===========================================================================
# STEP 5 — Recursive splitting (ported from the notebook)
# ===========================================================================
def process_and_unify_chunks(all_chunks: list[dict]) -> list[dict]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=INGEST_CHUNK_SIZE,
        chunk_overlap=INGEST_CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    unified: list[dict] = []
    for chunk in all_chunks:
        meta = chunk["metadata"]
        hierarchy = meta.get("Hierarchy", "")
        prefix = f"**[{hierarchy}]**\n" if hierarchy else ""
        category = chunk["category"]

        if category == "contains_image":
            chunk["content_processed"] = f"{prefix}{chunk.get('content_processed', chunk['content'])}"
            unified.append(chunk)
        else:
            text_to_split = chunk.get("content_processed", chunk["content"])
            for split_text in splitter.split_text(text_to_split):
                unified.append({
                    "content": chunk["content"],
                    "metadata": meta.copy(),
                    "category": category,
                    "content_processed": f"{prefix}{split_text.strip()}",
                })
    return unified


# ===========================================================================
# Chroma helpers
# ===========================================================================
def open_vector_db(embeddings) -> Chroma:
    return Chroma(
        persist_directory=VECTOR_DB_PATH,
        embedding_function=embeddings,
        collection_name=COLLECTION_NAME,
    )


def delete_source(vector_db: Chroma, source_name: str) -> int:
    """Remove all chunks belonging to a source. Returns how many were removed."""
    existing = vector_db.get(where={"source": source_name})
    ids = existing.get("ids", []) if existing else []
    if ids:
        vector_db.delete(ids=ids)
    return len(ids)


# ===========================================================================
# Orchestration
# ===========================================================================
def ingest_file(
    filename: str,
    data: bytes,
    embeddings,
    *,
    source_name: Optional[str] = None,
    use_vision: bool = True,
    use_table_summary: bool = True,
    use_keywords: bool = True,
    log: Logger = _noop,
) -> dict:
    """
    Full ingestion for a single uploaded file. Returns a stats dict:
        {source, chunks, images, tables, replaced, enriched}
    Raises on unrecoverable errors (bad file, embedding failure).
    """
    source = source_name or sanitize_source_name(filename)
    source_dir = Path(UPLOADS_DIR) / source
    source_dir.mkdir(parents=True, exist_ok=True)

    # 1 — save original
    original = source_dir / filename
    original.write_bytes(data)
    log(f"Saved upload → {original.name}")

    # 2 — convert to markdown (+ assets)
    content_path = convert_to_markdown(original, source_dir, log)

    # 3 — parse
    doc_chunks = parse_markdown_to_chunks(str(content_path), source)
    log(f"Parsed into {len(doc_chunks)} raw chunks.")

    # Decide what enrichment is actually possible right now.
    status = ollama_status()
    do_vision = use_vision and status["has_vision"]
    do_tables = use_table_summary and status["has_llm"]
    do_keywords = use_keywords and status["has_llm"]
    if (use_vision and not do_vision) or (use_table_summary and not do_tables):
        log("⚠ Ollama/model not available — skipping AI enrichment, embedding raw text instead.")

    image_chunks = [c for c in doc_chunks if c["category"] == "contains_image"]
    table_chunks = [c for c in doc_chunks if c["category"] == "contains_table"]

    # 4a — images
    enriched = 0
    if image_chunks:
        log(f"Describing {len(image_chunks)} image(s)…" if do_vision else f"Keeping {len(image_chunks)} image ref(s) as-is.")
        for i, chunk in enumerate(image_chunks, 1):
            if do_vision:
                try:
                    chunk["content_processed"] = summarize_image_with_vision(chunk, source_dir)
                    enriched += 1
                    log(f"  image {i}/{len(image_chunks)} ✓")
                except Exception as exc:
                    log(f"  image {i} skipped ({exc})")
                    chunk["content_processed"] = chunk["content"]
            else:
                # still record the asset name so traceability works
                m = re.search(r"!\[.*?\]\((.*?)\)", chunk["content"])
                if m:
                    chunk["metadata"]["asset"] = Path(m.group(1)).name
                chunk["content_processed"] = chunk["content"]

    # 4b — tables
    if table_chunks:
        log(f"Summarising {len(table_chunks)} table(s)…" if do_tables else f"Keeping {len(table_chunks)} table(s) as raw markdown.")
        for i, chunk in enumerate(table_chunks, 1):
            cleaned = clean_markdown_table(chunk["content"])
            if do_tables:
                try:
                    chunk["content_processed"] = summarize_table_with_qwen(cleaned)
                    enriched += 1
                    log(f"  table {i}/{len(table_chunks)} ✓")
                except Exception as exc:
                    log(f"  table {i} skipped ({exc})")
                    chunk["content_processed"] = cleaned
            else:
                chunk["content_processed"] = cleaned

    # 5 — split
    unified = process_and_unify_chunks(doc_chunks)
    log(f"Split into {len(unified)} embeddable chunks.")

    # 6 — keywords
    if do_keywords and unified:
        log("Extracting keywords…")
        for chunk in unified:
            try:
                chunk["keywords"] = extract_keywords_with_llm(chunk.get("content_processed", ""))
            except Exception:
                chunk["keywords"] = ""

    # 7 — embed (overwrite any prior chunks for this source)
    vector_db = open_vector_db(embeddings)
    replaced = delete_source(vector_db, source)
    if replaced:
        log(f"Replaced {replaced} existing chunk(s) for '{source}'.")

    documents = [
        Document(
            page_content=chunk["content_processed"],
            metadata={
                "source": str(chunk["metadata"].get("source", source)),
                "hierarchy": str(chunk["metadata"].get("Hierarchy", "")),
                "asset": str(chunk["metadata"].get("asset", "None")),
                "category": str(chunk["category"]),
                "keywords": str(chunk.get("keywords", "")),
                "original_content": str(chunk["content"]),
            },
        )
        for chunk in unified
    ]
    if documents:
        vector_db.add_documents(documents=documents)
    log(f"✓ Embedded {len(documents)} chunk(s) for '{source}'.")

    return {
        "source": source,
        "chunks": len(documents),
        "images": len(image_chunks),
        "tables": len(table_chunks),
        "replaced": replaced,
        "enriched": enriched,
    }
